"""File parsing service for extracting content from PDFs and DOCX files"""
import pdfplumber
from docx import Document
from typing import Optional
import os


def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract all text from a PDF file.
    
    Args:
        pdf_path: Path to the PDF file
        
    Returns:
        Extracted text content
    """
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    return text.strip()


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract all text from a DOCX file.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Extracted text content
    """
    text = ""
    try:
        doc = Document(docx_path)
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " | "
                text += "\n"
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    return text.strip()


def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from either a PDF or DOCX file based on file extension.
    
    Args:
        file_path: Path to the file (PDF or DOCX)
        
    Returns:
        Extracted text content
    """
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}. Only PDF and DOCX files are supported.")


def extract_tables_from_pdf(pdf_path: str) -> list:
    """
    Extract tables from a PDF file (useful for syllabus tables).
    
    Args:
        pdf_path: Path to the PDF file
        
    Returns:
        List of extracted tables
    """
    tables = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_tables = page.extract_tables()
                if page_tables:
                    tables.extend(page_tables)
    except Exception as e:
        raise ValueError(f"Failed to extract tables from PDF: {str(e)}")
    
    return tables
